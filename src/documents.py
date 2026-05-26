from __future__ import annotations

import hashlib
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader

import config

ENCODINGS = ("utf-8", "utf-8-sig", "latin-1", "cp1252")
BINARY_SKIP_SUFFIXES = {
    ".exe",
    ".dll",
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".zip",
    ".bin",
    ".doc",
}


@dataclass(frozen=True)
class TextChunk:
    text: str
    source_path: str
    start_line: int
    end_line: int
    chunk_index: int

    @property
    def chunk_id(self) -> str:
        digest = hashlib.sha256(
            f"{self.source_path}:{self.chunk_index}:{self.start_line}:{self.end_line}".encode()
        ).hexdigest()[:24]
        return f"chunk_{digest}"


def _decode_text(raw: bytes) -> str | None:
    for encoding in ENCODINGS:
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return None


def _normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _is_probably_text(path: Path) -> bool:
    suffix = path.suffix.lower()
    if suffix in config.OFFICE_EXTENSIONS:
        return False
    if suffix in config.TEXT_EXTENSIONS:
        return True
    if suffix == "" or suffix not in BINARY_SKIP_SUFFIXES:
        return True
    return False


def read_pdf_file(path: Path) -> str | None:
    try:
        reader = PdfReader(str(path))
    except Exception:
        return None

    pages: list[str] = []
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            continue
        text = _normalize_newlines(text).strip()
        if text:
            pages.append(f"--- Page {page_num} ---\n{text}")

    if not pages:
        return None
    return "\n\n".join(pages)


def read_docx_file(path: Path) -> str | None:
    try:
        document = Document(str(path))
    except Exception:
        return None

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    if not lines:
        return None
    return _normalize_newlines("\n".join(lines))


def read_text_file(path: Path) -> str | None:
    if not _is_probably_text(path):
        return None
    try:
        raw = path.read_bytes()
    except OSError:
        return None
    if not raw or b"\x00" in raw[:8192]:
        return None
    text = _decode_text(raw)
    if text is None:
        return None
    if not text.strip():
        return None
    return _normalize_newlines(text)


def read_document_file(path: Path) -> str | None:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    return read_text_file(path)


def iter_document_paths(root: Path) -> list[Path]:
    root = root.resolve()
    paths: list[Path] = []
    if root.is_file():
        return [root]

    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if any(part in config.SKIP_DIR_NAMES for part in path.parts):
            continue
        if path.name.startswith("."):
            continue
        paths.append(path)
    return sorted(paths)


def chunk_text(
    text: str,
    source_path: str,
    chunk_size: int,
    chunk_overlap: int,
) -> list[TextChunk]:
    lines = text.split("\n")
    chunks: list[TextChunk] = []
    start = 0
    chunk_index = 0

    while start < len(lines):
        current: list[str] = []
        size = 0
        line_idx = start

        while line_idx < len(lines):
            line = lines[line_idx]
            addition = len(line) + (1 if current else 0)
            if current and size + addition > chunk_size:
                break
            current.append(line)
            size += addition
            line_idx += 1
            if size >= chunk_size:
                break

        if not current:
            break

        end_line = start + len(current)
        chunk_body = "\n".join(current).strip()
        if chunk_body:
            chunks.append(
                TextChunk(
                    text=chunk_body,
                    source_path=source_path,
                    start_line=start + 1,
                    end_line=end_line,
                    chunk_index=chunk_index,
                )
            )
            chunk_index += 1

        if line_idx >= len(lines):
            break

        overlap_lines = max(1, chunk_overlap // max(1, (size // max(1, len(current)))))
        start = max(start + 1, end_line - overlap_lines)

    return chunks


def load_chunks_from_path(
    path: Path,
    chunk_size: int,
    chunk_overlap: int,
) -> list[TextChunk]:
    text = read_document_file(path)
    if text is None:
        return []
    rel = str(path.resolve())
    return chunk_text(text, rel, chunk_size, chunk_overlap)


def load_chunks_from_root(
    root: Path,
    chunk_size: int,
    chunk_overlap: int,
) -> list[TextChunk]:
    all_chunks: list[TextChunk] = []
    for doc_path in iter_document_paths(root):
        all_chunks.extend(load_chunks_from_path(doc_path, chunk_size, chunk_overlap))
    return all_chunks
